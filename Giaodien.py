import io
import re
from docx import Document
from PyQt5 import QtCore, QtGui, QtWidgets
from PyQt5.QtWidgets import QFileDialog
from Elgamal import gcd, binhPhuong, kiemTraNguyenTo, thucHienKy, kiemTraChuKy
import random

class Ui_MainWindow(object):
    def setupUi(self, MainWindow):
        MainWindow.setObjectName("MainWindow")
        MainWindow.resize(1421, 894)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Preferred, QtWidgets.QSizePolicy.Preferred)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(MainWindow.sizePolicy().hasHeightForWidth())
        MainWindow.setSizePolicy(sizePolicy)
        self.centralwidget = QtWidgets.QWidget(MainWindow)
        self.centralwidget.setObjectName("centralwidget")
        self.label = QtWidgets.QLabel(self.centralwidget)
        self.label.setGeometry(QtCore.QRect(540, 10, 337, 37))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(20)
        self.label.setFont(font)
        self.label.setObjectName("label")
        self.groupBox = QtWidgets.QGroupBox(self.centralwidget)
        self.groupBox.setGeometry(QtCore.QRect(30, 60, 1361, 201))
        self.groupBox.setTitle("")
        self.groupBox.setObjectName("groupBox")
        self.label_2 = QtWidgets.QLabel(self.groupBox)
        self.label_2.setGeometry(QtCore.QRect(10, 10, 243, 29))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(15)
        font.setBold(False)
        font.setItalic(False)
        font.setUnderline(False)
        font.setWeight(50)
        font.setStrikeOut(False)
        self.label_2.setFont(font)
        self.label_2.setTextFormat(QtCore.Qt.PlainText)
        self.label_2.setObjectName("label_2")
        self.label_3 = QtWidgets.QLabel(self.groupBox)
        self.label_3.setGeometry(QtCore.QRect(50, 50, 147, 26))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(14)
        self.label_3.setFont(font)
        self.label_3.setObjectName("label_3")
        self.label_4 = QtWidgets.QLabel(self.groupBox)
        self.label_4.setGeometry(QtCore.QRect(140, 100, 49, 26))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(14)
        self.label_4.setFont(font)
        self.label_4.setObjectName("label_4")
        self.label_5 = QtWidgets.QLabel(self.groupBox)
        self.label_5.setGeometry(QtCore.QRect(50, 150, 141, 26))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(14)
        self.label_5.setFont(font)
        self.label_5.setObjectName("label_5")
        self.label_6 = QtWidgets.QLabel(self.groupBox)
        self.label_6.setGeometry(QtCore.QRect(480, 50, 47, 26))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(14)
        self.label_6.setFont(font)
        self.label_6.setObjectName("label_6")
        self.label_7 = QtWidgets.QLabel(self.groupBox)
        self.label_7.setGeometry(QtCore.QRect(1010, 40, 158, 26))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(14)
        self.label_7.setFont(font)
        self.label_7.setObjectName("label_7")
        self.label_8 = QtWidgets.QLabel(self.groupBox)
        self.label_8.setGeometry(QtCore.QRect(1030, 90, 141, 26))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(14)
        self.label_8.setFont(font)
        self.label_8.setObjectName("label_8")
        self.soNguyenToP = QtWidgets.QLineEdit(self.groupBox)
        self.soNguyenToP.setGeometry(QtCore.QRect(210, 50, 161, 31))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(14)
        self.soNguyenToP.setFont(font)
        self.soNguyenToP.setObjectName("soNguyenToP")
        self.alpha = QtWidgets.QLineEdit(self.groupBox)
        self.alpha.setGeometry(QtCore.QRect(210, 100, 161, 31))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(14)
        self.alpha.setFont(font)
        self.alpha.setObjectName("alpha")
        self.beta = QtWidgets.QLineEdit(self.groupBox)
        self.beta.setGeometry(QtCore.QRect(210, 150, 161, 31))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(14)
        self.beta.setFont(font)
        self.beta.setObjectName("beta")
        self.label_9 = QtWidgets.QLabel(self.groupBox)
        self.label_9.setGeometry(QtCore.QRect(460, 10, 158, 29))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(15)
        font.setUnderline(False)
        font.setStrikeOut(False)
        self.label_9.setFont(font)
        self.label_9.setTextFormat(QtCore.Qt.PlainText)
        self.label_9.setObjectName("label_9")
        self.soA = QtWidgets.QLineEdit(self.groupBox)
        self.soA.setGeometry(QtCore.QRect(540, 50, 161, 31))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(14)
        self.soA.setFont(font)
        self.soA.setObjectName("soA")
        self.soK = QtWidgets.QLineEdit(self.groupBox)
        self.soK.setGeometry(QtCore.QRect(1180, 40, 161, 31))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(14)
        self.soK.setFont(font)
        self.soK.setObjectName("soK")
        self.gamma = QtWidgets.QLineEdit(self.groupBox)
        self.gamma.setGeometry(QtCore.QRect(1180, 90, 161, 31))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(14)
        self.gamma.setFont(font)
        self.gamma.setObjectName("gamma")
        self.lamMoi = QtWidgets.QPushButton(self.groupBox)
        self.lamMoi.setGeometry(QtCore.QRect(810, 140, 121, 38))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(15)
        self.lamMoi.setFont(font)
        self.lamMoi.setObjectName("lamMoi")
        self.taoKhoaNgauNhien = QtWidgets.QPushButton(self.groupBox)
        self.taoKhoaNgauNhien.setGeometry(QtCore.QRect(750, 80, 231, 38))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(15)
        self.taoKhoaNgauNhien.setFont(font)
        self.taoKhoaNgauNhien.setObjectName("taoKhoaNgauNhien")
        self.tinh = QtWidgets.QPushButton(self.groupBox)
        self.tinh.setGeometry(QtCore.QRect(830, 20, 81, 38))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(15)
        self.tinh.setFont(font)
        self.tinh.setObjectName("tinh")
        self.groupBox_2 = QtWidgets.QGroupBox(self.centralwidget)
        self.groupBox_2.setGeometry(QtCore.QRect(30, 300, 651, 521))
        self.groupBox_2.setTitle("")
        self.groupBox_2.setObjectName("groupBox_2")
        self.label_10 = QtWidgets.QLabel(self.groupBox_2)
        self.label_10.setGeometry(QtCore.QRect(220, 10, 168, 29))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(15)
        font.setUnderline(False)
        font.setStrikeOut(False)
        self.label_10.setFont(font)
        self.label_10.setTextFormat(QtCore.Qt.PlainText)
        self.label_10.setObjectName("label_10")
        self.label_12 = QtWidgets.QLabel(self.groupBox_2)
        self.label_12.setGeometry(QtCore.QRect(20, 60, 114, 26))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(14)
        self.label_12.setFont(font)
        self.label_12.setObjectName("label_12")
        self.label_13 = QtWidgets.QLabel(self.groupBox_2)
        self.label_13.setGeometry(QtCore.QRect(50, 310, 75, 26))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(14)
        self.label_13.setFont(font)
        self.label_13.setObjectName("label_13")
        self.file = QtWidgets.QPushButton(self.groupBox_2)
        self.file.setGeometry(QtCore.QRect(510, 100, 93, 35))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(14)
        self.file.setFont(font)
        self.file.setObjectName("file")
        self.ky = QtWidgets.QPushButton(self.groupBox_2)
        self.ky.setGeometry(QtCore.QRect(270, 260, 93, 35))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(14)
        self.ky.setFont(font)
        self.ky.setObjectName("ky")
        self.chuyen = QtWidgets.QPushButton(self.groupBox_2)
        self.chuyen.setGeometry(QtCore.QRect(510, 330, 93, 35))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(14)
        self.chuyen.setFont(font)
        self.chuyen.setObjectName("chuyen")
        self.luu = QtWidgets.QPushButton(self.groupBox_2)
        self.luu.setGeometry(QtCore.QRect(510, 400, 93, 35))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(14)
        self.luu.setFont(font)
        self.luu.setObjectName("luu")
        self.vanBanKy1 = QtWidgets.QTextEdit(self.groupBox_2)
        self.vanBanKy1.setGeometry(QtCore.QRect(170, 60, 291, 191))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(14)
        self.vanBanKy1.setFont(font)
        self.vanBanKy1.setObjectName("vanBanKy1")
        self.chuKy1 = QtWidgets.QTextEdit(self.groupBox_2)
        self.chuKy1.setGeometry(QtCore.QRect(170, 310, 291, 191))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(14)
        self.chuKy1.setFont(font)
        self.chuKy1.setObjectName("chuKy1")
        self.groupBox_3 = QtWidgets.QGroupBox(self.centralwidget)
        self.groupBox_3.setGeometry(QtCore.QRect(740, 300, 651, 521))
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Expanding, QtWidgets.QSizePolicy.Preferred)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.groupBox_3.sizePolicy().hasHeightForWidth())
        self.groupBox_3.setSizePolicy(sizePolicy)
        self.groupBox_3.setTitle("")
        self.groupBox_3.setObjectName("groupBox_3")
        self.label_11 = QtWidgets.QLabel(self.groupBox_3)
        self.label_11.setGeometry(QtCore.QRect(240, 10, 162, 29))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(15)
        font.setUnderline(False)
        font.setStrikeOut(False)
        self.label_11.setFont(font)
        self.label_11.setTextFormat(QtCore.Qt.PlainText)
        self.label_11.setObjectName("label_11")
        self.label_14 = QtWidgets.QLabel(self.groupBox_3)
        self.label_14.setGeometry(QtCore.QRect(30, 50, 114, 26))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(14)
        self.label_14.setFont(font)
        self.label_14.setObjectName("label_14")
        self.label_15 = QtWidgets.QLabel(self.groupBox_3)
        self.label_15.setGeometry(QtCore.QRect(70, 190, 75, 26))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(14)
        self.label_15.setFont(font)
        self.label_15.setObjectName("label_15")
        self.fileVanBan = QtWidgets.QPushButton(self.groupBox_3)
        self.fileVanBan.setGeometry(QtCore.QRect(510, 80, 124, 35))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(14)
        self.fileVanBan.setFont(font)
        self.fileVanBan.setObjectName("fileVanBan")
        self.fileChuKy = QtWidgets.QPushButton(self.groupBox_3)
        self.fileChuKy.setGeometry(QtCore.QRect(510, 220, 114, 35))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(14)
        self.fileChuKy.setFont(font)
        self.fileChuKy.setObjectName("fileChuKy")
        self.kiemTraChuKy = QtWidgets.QPushButton(self.groupBox_3)
        self.kiemTraChuKy.setGeometry(QtCore.QRect(240, 340, 159, 35))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(14)
        self.kiemTraChuKy.setFont(font)
        self.kiemTraChuKy.setObjectName("kiemTraChuKy")
        self.label_16 = QtWidgets.QLabel(self.groupBox_3)
        self.label_16.setGeometry(QtCore.QRect(40, 390, 107, 26))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(14)
        self.label_16.setFont(font)
        self.label_16.setObjectName("label_16")
        self.vanBanKy2 = QtWidgets.QTextEdit(self.groupBox_3)
        self.vanBanKy2.setGeometry(QtCore.QRect(170, 50, 311, 111))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(14)
        self.vanBanKy2.setFont(font)
        self.vanBanKy2.setObjectName("vanBanKy2")
        self.chuKy2 = QtWidgets.QTextEdit(self.groupBox_3)
        self.chuKy2.setGeometry(QtCore.QRect(170, 190, 311, 131))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(14)
        self.chuKy2.setFont(font)
        self.chuKy2.setObjectName("chuKy2")
        self.thongBao = QtWidgets.QTextEdit(self.groupBox_3)
        self.thongBao.setGeometry(QtCore.QRect(170, 390, 311, 111))
        font = QtGui.QFont()
        font.setFamily("Times New Roman")
        font.setPointSize(14)
        self.thongBao.setFont(font)
        self.thongBao.setObjectName("thongBao")
        MainWindow.setCentralWidget(self.centralwidget)
        self.menubar = QtWidgets.QMenuBar(MainWindow)
        self.menubar.setGeometry(QtCore.QRect(0, 0, 1421, 26))
        self.menubar.setObjectName("menubar")
        MainWindow.setMenuBar(self.menubar)
        self.statusbar = QtWidgets.QStatusBar(MainWindow)
        self.statusbar.setObjectName("statusbar")
        MainWindow.setStatusBar(self.statusbar)

        self.retranslateUi(MainWindow)
        QtCore.QMetaObject.connectSlotsByName(MainWindow)

        self.taoKhoaNgauNhien.clicked.connect(self.taoGiaTri)
        self.lamMoi.clicked.connect(self.xoa)
        self.tinh.clicked.connect(self.tinhToan)
        self.chuyen.clicked.connect(self.chuyenVB)
        self.ky.clicked.connect(self.kyVB)
        self.kiemTraChuKy.clicked.connect(self.ktKy)
        self.file.clicked.connect(self.file1)
        self.luu.clicked.connect(self.luuFile)
        self.fileVanBan.clicked.connect(self.file2)
        self.fileChuKy.clicked.connect(self.file3)

    def retranslateUi(self, MainWindow):
        _translate = QtCore.QCoreApplication.translate
        MainWindow.setWindowTitle(_translate("MainWindow", "Chữ ký Elgamal"))
        self.label.setText(_translate("MainWindow", "Chữ Ký Điện Tử Elgamal"))
        self.label_2.setText(_translate("MainWindow", "Khóa công khai (p, α, β)"))
        self.label_3.setText(_translate("MainWindow", "Số nguyên tố p:"))
        self.label_4.setText(_translate("MainWindow", "Số α:"))
        self.label_5.setText(_translate("MainWindow", "β = α^a mod p:"))
        self.label_6.setText(_translate("MainWindow", "Số a:"))
        self.label_7.setText(_translate("MainWindow", "Số ngẫu nhiên k:"))
        self.label_8.setText(_translate("MainWindow", "γ = α^k mod p:"))
        self.label_9.setText(_translate("MainWindow", "Khóa bí mật (a)"))
        self.lamMoi.setText(_translate("MainWindow", "Làm mới"))
        self.taoKhoaNgauNhien.setText(_translate("MainWindow", "Tạo khóa ngẫu nhiên"))
        self.tinh.setText(_translate("MainWindow", "Tính"))
        self.label_10.setText(_translate("MainWindow", "Phát sinh chữ ký"))
        self.label_12.setText(_translate("MainWindow", "Văn bản ký:"))
        self.label_13.setText(_translate("MainWindow", "Chữ ký:"))
        self.file.setText(_translate("MainWindow", "File"))
        self.ky.setText(_translate("MainWindow", "Ký"))
        self.chuyen.setText(_translate("MainWindow", "Chuyển"))
        self.luu.setText(_translate("MainWindow", "Lưu"))
        self.label_11.setText(_translate("MainWindow", "Kiểm tra chữ ký"))
        self.label_14.setText(_translate("MainWindow", "Văn bản ký:"))
        self.label_15.setText(_translate("MainWindow", "Chữ ký:"))
        self.fileVanBan.setText(_translate("MainWindow", "File văn bản"))
        self.fileChuKy.setText(_translate("MainWindow", "File chữ ký"))
        self.kiemTraChuKy.setText(_translate("MainWindow", "Kiểm tra chữ ký"))
        self.label_16.setText(_translate("MainWindow", "Thông báo:"))

    def taoGiaTri(self):
        p = random.randint(100, 10 ** 4)
        while not kiemTraNguyenTo(p) == 1:
            p = random.randint(100, 10 ** 4)
        self.soNguyenToP.setText(str(p))
        alpha = random.randint(1, p - 1)
        self.alpha.setText(str(alpha))
        a = random.randint(2, p - 2)
        self.soA.setText(str(a))
        beta = binhPhuong(alpha, a, p)
        self.beta.setText(str(beta))
        k = 0
        while not gcd(k, p - 1) == 1:
            k = random.randint(1, p - 2)
        self.soK.setText(str(k))
        gamma = binhPhuong(alpha, k, p)
        self.gamma.setText(str(gamma))

    def xoa(self):
        self.soNguyenToP.clear()
        self.alpha.clear()
        self.beta.clear()
        self.soK.clear()
        self.gamma.clear()
        self.soA.clear()
        self.vanBanKy1.clear()
        self.vanBanKy2.clear()
        self.chuKy1.clear()
        self.chuKy2.clear()
        self.thongBao.clear()
    def tinhToan(self):
        mes = QtWidgets.QMessageBox()
        a = self.soA.text()
        alpha = self.alpha.text()
        k = self.soK.text()
        p = self.soNguyenToP.text()
        if p == "" or a == "" or alpha == "" or k == "":
            mes.setInformativeText("Hãy nhập đầy đủ các giá trị cần tính!")
            mes.exec()
        elif kiemTraNguyenTo(int(p)) == 0 or int(alpha)<1 or int(alpha)>int(p)-1 or int(a)<2 \
        or int(a)>int(p)-2 or int(k)<1 or int(k)>int(p)-2 or gcd(int(k),int(p)-1) != 1:
            mes.setInformativeText("Hãy nhập lại!")
            mes.exec()
        else:
            beta = binhPhuong(int(alpha), int(a), int(p))
            self.beta.setText(str(beta))
            gamma = binhPhuong(int(alpha), int(k), int(p))
            self.gamma.setText(str(gamma))
    def chuyenVB(self):
        vb1 = self.vanBanKy1.toPlainText()
        self.vanBanKy2.setText(str(vb1))
        vk1 = self.chuKy1.toPlainText()
        self.chuKy2.setText(str(vk1))
    def kyVB(self):
        mes = QtWidgets.QMessageBox()
        try:
            vb1 = self.vanBanKy1.toPlainText()
            p = int(self.soNguyenToP.text())
            a = int(self.soA.text())
            k = int(self.soK.text())
            gamma = int(self.gamma.text())
            deta = thucHienKy(vb1,p,a,k,gamma)
            self.chuKy1.setText(str(gamma) + "," + str(deta))
        except Exception as e:
            mes.setInformativeText(f"Lỗi khi thực hiện: {str(e)}")
            mes.exec()
    def ktKy(self):
        mes = QtWidgets.QMessageBox()
        vb2 = self.vanBanKy2.toPlainText()
        p = int(self.soNguyenToP.text())
        ck2 = self.chuKy2.toPlainText()
        if vb2 == "" or ck2 == "":
            mes.setInformativeText("Hãy nhập đầy đủ thông tin!")
            mes.exec()
        elif not re.match(r'^\d+,\d+$', ck2):
            self.thongBao.setText("Chữ ký không chính xác!")
        else:
            chuoi = ck2.split(",")
            gamma = int(chuoi[0])
            deta = int(chuoi[1])
            kt = kiemTraChuKy(vb2, gamma, deta, int(self.alpha.text()), int(self.beta.text()), p)
            if kt == 1:
                self.thongBao.setText("Chữ ký đúng!")
            else :
                self.thongBao.setText("Chữ ký không chính xác!")
    def file1(self):
        filename = QFileDialog.getOpenFileName()
        file_path = filename[0]

        if file_path:
            if file_path.endswith(".txt"):
                with open(file_path, encoding='utf8') as file:
                    content = file.read()
                self.vanBanKy1.setText(str(content))
            elif file_path.endswith(".docx"):
                document = Document(file_path)
                content = []

                for paragraph in document.paragraphs:
                    runs = paragraph.runs
                    text = ""
                    for run in runs:
                        run_text = run.text
                        if run.bold:
                            run_text = f"<b>{run_text}</b>"
                        if run.italic:
                            run_text = f"<i>{run_text}</i>"
                        if run.underline:
                            run_text = f"<u>{run_text}</u>"
                        if run.font.name:
                            font = run.font.name
                            run_text = f'<span style="font-family: {font}">{run_text}</span>'
                        if run.font.size:
                            size = run.font.size.pt
                            run_text = f'<span style="font-size: {size}pt">{run_text}</span>'
                        if run.font.color.rgb:
                            color = run.font.color.rgb
                            run_text = f'<span style="color:#{color}">{run_text}</span>'
                        if run.font.highlight_color:
                            high_color = run.font.highlight_color
                            high_color = str(high_color)[:8].lower()
                            if high_color == "bright_g":
                                high_color = "#00FF00"
                            elif high_color == "turquois":
                                high_color = "#00FFFF"
                            elif high_color == "pink (5)":
                                high_color = "#FF00FF"
                            elif high_color == "blue (2)":
                                high_color = "#0000FF"
                            elif high_color == "red (6)":
                                high_color = "#FF0000"
                            elif high_color == "dark_blu":
                                high_color = "#000080"
                            elif high_color == "teal (10":
                                high_color = "#008080"
                            elif high_color == "green (1":
                                high_color = "#00FF00"
                            elif high_color == "violet (":
                                high_color = "#EE82EE"
                            elif high_color == "dark_red":
                                high_color = "#800000"
                            elif high_color == "dark_yel":
                                high_color = "#808000"
                            elif high_color == "gray_50 ":
                                high_color = "#808080"
                            elif high_color == "gray_25 ":
                                high_color = "#C0C0C0"
                            elif high_color == "black (1":
                                high_color = "#000000"
                            elif high_color == "yellow (":
                                high_color = "#FFFF00"
                            run_text = f'<mark style="background-color: {high_color};">{run_text}</mark>'
                        if run.font.superscript:
                            run_text = f'<sup>{run_text}</sup>'
                        if run.font.subscript:
                            run_text = f'<sub>{run_text}</sub>'
                        if run.font.strike:
                            run_text = f'<s>{run_text}</s>'
                        text += run_text
                    content.append(text)

                html_content = "<br>".join(content)
                self.vanBanKy1.setHtml(html_content)

            else:
                mes = QtWidgets.QMessageBox()
                mes.setInformativeText("File không hợp lệ")
                mes.exec()
    def file2(self):
        filename = QFileDialog.getOpenFileName()
        file_path = filename[0]

        if file_path:
            if file_path.endswith(".txt"):
                with open(file_path, encoding='utf8') as file:
                    content = file.read()
                self.vanBanKy2.setText(str(content))
            elif file_path.endswith(".docx"):
                document = Document(file_path)
                content = []

                for paragraph in document.paragraphs:
                    runs = paragraph.runs
                    text = ""
                    for run in runs:
                        run_text = run.text
                        if run.bold:
                            run_text = f"<b>{run_text}</b>"
                        if run.italic:
                            run_text = f"<i>{run_text}</i>"
                        if run.underline:
                            run_text = f"<u>{run_text}</u>"
                        if run.font.name:
                            font = run.font.name
                            run_text = f'<span style="font-family: {font}">{run_text}</span>'
                        if run.font.size:
                            size = run.font.size.pt
                            run_text = f'<span style="font-size: {size}pt">{run_text}</span>'
                        if run.font.color.rgb:
                            color = run.font.color.rgb
                            run_text = f'<span style="color:#{color}">{run_text}</span>'
                        if run.font.highlight_color:
                            high_color = run.font.highlight_color
                            high_color = str(high_color)[:8].lower()
                            if high_color == "bright_g":
                                high_color = "#00FF00"
                            elif high_color == "turquois":
                                high_color = "#00FFFF"
                            elif high_color == "pink (5)":
                                high_color = "#FF00FF"
                            elif high_color == "blue (2)":
                                high_color = "#0000FF"
                            elif high_color == "red (6)":
                                high_color = "#FF0000"
                            elif high_color == "dark_blu":
                                high_color = "#000080"
                            elif high_color == "teal (10":
                                high_color = "#008080"
                            elif high_color == "green (1":
                                high_color = "#00FF00"
                            elif high_color == "violet (":
                                high_color = "#EE82EE"
                            elif high_color == "dark_red":
                                high_color = "#800000"
                            elif high_color == "dark_yel":
                                high_color = "#808000"
                            elif high_color == "gray_50 ":
                                high_color = "#808080"
                            elif high_color == "gray_25 ":
                                high_color = "#C0C0C0"
                            elif high_color == "black (1":
                                high_color = "#000000"
                            elif high_color == "yellow (":
                                high_color = "#FFFF00"
                            run_text = f'<mark style="background-color: {high_color};">{run_text}</mark>'
                        if run.font.superscript:
                            run_text = f'<sup>{run_text}</sup>'
                        if run.font.subscript:
                            run_text = f'<sub>{run_text}</sub>'
                        if run.font.strike:
                            run_text = f'<s>{run_text}</s>'
                        text += run_text
                    content.append(text)

                html_content = "<br>".join(content)
                self.vanBanKy2.setHtml(html_content)

            else:
                mes = QtWidgets.QMessageBox()
                mes.setInformativeText("File không hợp lệ")
                mes.exec()
    def file3(self):
        filename = QFileDialog.getOpenFileName()
        file_path = filename[0]

        if file_path:
            if file_path.endswith(".txt"):
                with open(file_path, encoding='utf8') as file:
                    content = file.read()
                self.chuKy2.setText(str(content))
            elif file_path.endswith(".docx"):
                document = Document(file_path)
                content = []

                for paragraph in document.paragraphs:
                    runs = paragraph.runs
                    text = ""
                    for run in runs:
                        run_text = run.text
                        if run.bold:
                            run_text = f"<b>{run_text}</b>"
                        if run.italic:
                            run_text = f"<i>{run_text}</i>"
                        if run.underline:
                            run_text = f"<u>{run_text}</u>"
                        if run.font.name:
                            font = run.font.name
                            run_text = f'<span style="font-family: {font}">{run_text}</span>'
                        if run.font.size:
                            size = run.font.size.pt
                            run_text = f'<span style="font-size: {size}pt">{run_text}</span>'
                        if run.font.color.rgb:
                            color = run.font.color.rgb
                            run_text = f'<span style="color:#{color}">{run_text}</span>'
                        if run.font.highlight_color:
                            high_color = run.font.highlight_color
                            high_color = str(high_color)[:8].lower()
                            if high_color == "bright_g":
                                high_color = "#00FF00"
                            elif high_color == "turquois":
                                high_color = "#00FFFF"
                            elif high_color == "pink (5)":
                                high_color = "#FF00FF"
                            elif high_color == "blue (2)":
                                high_color = "#0000FF"
                            elif high_color == "red (6)":
                                high_color = "#FF0000"
                            elif high_color == "dark_blu":
                                high_color = "#000080"
                            elif high_color == "teal (10":
                                high_color = "#008080"
                            elif high_color == "green (1":
                                high_color = "#00FF00"
                            elif high_color == "violet (":
                                high_color = "#EE82EE"
                            elif high_color == "dark_red":
                                high_color = "#800000"
                            elif high_color == "dark_yel":
                                high_color = "#808000"
                            elif high_color == "gray_50 ":
                                high_color = "#808080"
                            elif high_color == "gray_25 ":
                                high_color = "#C0C0C0"
                            elif high_color == "black (1":
                                high_color = "#000000"
                            elif high_color == "yellow (":
                                high_color = "#FFFF00"
                            run_text = f'<mark style="background-color: {high_color};">{run_text}</mark>'
                        if run.font.superscript:
                            run_text = f'<sup>{run_text}</sup>'
                        if run.font.subscript:
                            run_text = f'<sub>{run_text}</sub>'
                        if run.font.strike:
                            run_text = f'<s>{run_text}</s>'
                        text += run_text
                    content.append(text)

                html_content = "<br>".join(content)
                self.chuKy2.setHtml(html_content)

            else:
                mes = QtWidgets.QMessageBox()
                mes.setInformativeText("File không hợp lệ")
                mes.exec()
    def luuFile(self):
        dialog = QFileDialog()
        dialog.setAcceptMode(QFileDialog.AcceptSave)
        dialog.setDefaultSuffix("txt")
        dialog.setNameFilter("Tệp tin văn bản (*.txt);;Microsoft Word (*.docx)")

        if dialog.exec_():
            file_path = dialog.selectedFiles()[0]
            content = self.chuKy1.toPlainText()
            mes = QtWidgets.QMessageBox()

            if file_path.endswith(".txt"):
                try:
                    with io.open(file_path, 'w', encoding='utf8') as file:
                        file.write(content)
                    mes.setInformativeText("Lưu tệp tin thành công")
                except Exception as e:
                    mes.setInformativeText(f"Lỗi khi lưu tệp tin: {str(e)}")

            elif file_path.endswith(".docx"):
                try:
                    document = Document()
                    document.add_paragraph(content)
                    document.save(file_path)
                    mes.setInformativeText("Lưu tệp tin thành công")
                except Exception as e:
                    mes.setInformativeText(f"Lỗi khi lưu tệp tin: {str(e)}")

            else:
                mes.setInformativeText("Định dạng tệp tin không hợp lệ")
            mes.exec()

if __name__ == "__main__":
    import sys
    app = QtWidgets.QApplication(sys.argv)
    MainWindow = QtWidgets.QMainWindow()
    ui = Ui_MainWindow()
    ui.setupUi(MainWindow)
    MainWindow.show()
    sys.exit(app.exec_())
